

import itertools
import json
import logging
import re
from collections import OrderedDict, defaultdict
from dataclasses import dataclass, field
from datetime import date, datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import fitz  # PyMuPDF
import phonenumbers
import pycountry
from dateutil import parser as dateutil_parser
from dateutil.relativedelta import relativedelta
from django.db import transaction
from docx import Document as DocxDocument
from nameparser import HumanName
from rapidfuzz import fuzz
from rapidfuzz import process as rf_process

try:
    import dateparser

    HAS_DATEPARSER = True
except ImportError:  # pragma: no cover - optional dependency
    HAS_DATEPARSER = False

try:
    import spacy

    HAS_SPACY = True
except ImportError:  # pragma: no cover
    HAS_SPACY = False

try:
    from gliner import GLiNER

    HAS_GLINER = True
except ImportError:  # pragma: no cover
    HAS_GLINER = False

from .models import (
    Candidate,
    EmploymentType,
    LanguageProficiency,
    ProficiencyLevel,
    Resume,
    ResumeAchievement,
    ResumeCertification,
    ResumeCustomSection,
    ResumeEducation,
    ResumeLanguage,
    ResumePersonalInformation,
    ResumeProject,
    ResumeSkill,
    ResumeStatus,
    ResumeWorkExperience,
    SkillCategory,
)

logger = logging.getLogger(__name__)


# ======================================================================
# Lazy singleton model loaders
# ======================================================================

_SPACY_NLP = None
_GLINER_MODEL = None

GLINER_LABELS = [
    "Full Name",
    "Job Title",
    "Company",
    "Skill",
    "Programming Language",
    "Framework",
    "Database",
    "Cloud Platform",
    "Tool",
    "Degree",
    "College",
    "University",
    "Certification",
    "Project",
    "Technology",
    "Achievement",
    "Responsibility",
    "Nationality",
    "Address",
]


def get_spacy_model():
    
    global _SPACY_NLP
    if not HAS_SPACY:
        return None
    if _SPACY_NLP is None:
        for model_name in ("en_core_web_trf", "en_core_web_lg", "en_core_web_sm"):
            try:
                _SPACY_NLP = spacy.load(model_name)
                break
            except OSError:
                continue
        if _SPACY_NLP is None:
            logger.warning("No spaCy model could be loaded; spaCy NER disabled.")
    return _SPACY_NLP


def get_gliner_model():
    
    global _GLINER_MODEL
    if not HAS_GLINER:
        return None
    if _GLINER_MODEL is None:
        try:
            _GLINER_MODEL = GLiNER.from_pretrained("urchade/gliner_multi-v2.1")
        except Exception:  # pragma: no cover - offline / missing weights
            logger.warning("GLiNER model could not be loaded; GLiNER NER disabled.")
            _GLINER_MODEL = None
    return _GLINER_MODEL


# ======================================================================
# Confidence-scored value container
# ======================================================================

@dataclass
class ScoredValue:
   

    value: Any
    confidence: float
    source: str


class FieldResolver:
   

    def __init__(self, name: str):
        self.name = name
        self.candidates: List[ScoredValue] = []

    def add(self, value: Any, confidence: float, source: str):
        if value in (None, "", [], {}):
            return
        self.candidates.append(ScoredValue(value=value, confidence=confidence, source=source))

    @property
    def best(self) -> Optional[ScoredValue]:
        if not self.candidates:
            return None
        return max(self.candidates, key=lambda c: c.confidence)

    @property
    def value(self) -> Any:
        best = self.best
        return best.value if best else None

    def as_dict(self) -> Dict[str, Any]:
        return {
            "value": self.value,
            "confidence": self.best.confidence if self.best else 0.0,
            "alternatives": [
                {"value": c.value, "confidence": c.confidence, "source": c.source}
                for c in sorted(self.candidates, key=lambda c: -c.confidence)
            ],
        }


# ======================================================================
# Regex patterns
# ======================================================================

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")

URL_RE = re.compile(
    r"(?:https?://)?(?:www\.)?[a-zA-Z0-9\-]+(?:\.[a-zA-Z0-9\-]+)+(?:/[^\s,;()<>]*)?",
    re.IGNORECASE,
)

LINKEDIN_RE = re.compile(r"(?:https?://)?(?:[a-z]{2,3}\.)?linkedin\.com/[^\s,;()<>]+", re.I)
GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[^\s,;()<>]+", re.I)

PHONE_CANDIDATE_RE = re.compile(
    r"(\+?\d{1,3}[\s.\-]?)?(\(?\d{2,4}\)?[\s.\-]?)?\d{3,4}[\s.\-]?\d{3,4}(?:[\s.\-]?\d{2,4})?"
)

DATE_RANGE_SEP_RE = re.compile(
    r"\s*(?:-|–|—|to|until|through)\s*", re.IGNORECASE
)

PRESENT_RE = re.compile(r"^(present|current|till date|ongoing|now)$", re.IGNORECASE)

DATE_TOKEN_RE = re.compile(
    r"(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"[\s.,\-]*\d{0,2}[\s.,\-]*\d{2,4}"
    r"|\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}"
    r"|\d{4})",
    re.IGNORECASE,
)

CGPA_RE = re.compile(r"\bC?GPA\b\s*[:\-]?\s*(\d(?:\.\d{1,2})?)\s*(?:/\s*(\d(?:\.\d{1,2})?))?", re.I)
PERCENTAGE_RE = re.compile(r"(\d{1,3}(?:\.\d{1,2})?)\s*%")

BULLET_PREFIX_RE = re.compile(r"^[\s]*[\u2022\u25CF\u25E6\-\*\u2013\u2014>\u2023\u2043]+\s*")

POSTAL_CODE_RE = re.compile(r"\b\d{5,6}(?:-\d{4})?\b")

DEGREE_KEYWORDS = [
    "bachelor", "master", "b.tech", "btech", "b.e", "be ", "m.tech", "mtech",
    "b.sc", "bsc", "m.sc", "msc", "mba", "bba", "phd", "ph.d", "doctorate",
    "diploma", "associate degree", "b.com", "bcom", "m.com", "mcom", "bca",
    "mca", "llb", "ll.b", "high school", "hsc", "ssc", "matriculation",
    "post graduate", "postgraduate", "undergraduate", "b.a", "ba ", "m.a",
    "ma ", "b.arch", "m.arch",
]


# ======================================================================
# Section heading synonyms (fuzzy matched with RapidFuzz)
# ======================================================================

SECTION_HEADING_SYNONYMS: Dict[str, List[str]] = {
    "education": [
        "education", "academic qualification", "academic qualifications",
        "qualifications", "academic background", "educational background",
        "academic details", "academic profile",
    ],
    "skills": [
        "skills", "technical skills", "core skills", "competencies",
        "expertise", "tech stack", "key skills", "areas of expertise",
        "skill set", "core competencies",
    ],
    "experience": [
        "experience", "employment", "professional experience",
        "work history", "career history", "work experience",
        "employment history", "professional background",
        "relevant experience",
    ],
    "projects": [
        "projects", "personal projects", "academic projects",
        "professional projects", "key projects", "project experience",
        "notable projects",
    ],
    "certifications": [
        "certifications", "licenses", "credentials", "certificates",
        "professional certifications", "licenses and certifications",
    ],
    "languages": [
        "languages", "language proficiency", "languages known",
        "spoken languages",
    ],
    "achievements": [
        "achievements", "awards", "honors", "honours",
        "awards and achievements", "accomplishments", "recognitions",
    ],
    "summary": [
        "summary", "profile summary", "professional summary",
        "career objective", "objective", "about me", "profile",
        "executive summary",
    ],
}

# Flat lookup: synonym text -> canonical section key
_ALL_SYNONYMS: List[Tuple[str, str]] = [
    (syn, canon) for canon, syns in SECTION_HEADING_SYNONYMS.items() for syn in syns
]
_SYNONYM_TEXTS = [s for s, _ in _ALL_SYNONYMS]


# ======================================================================
# Skill dictionary
# ======================================================================

SKILL_DICTIONARY: Dict[str, List[str]] = {
    "PROGRAMMING_LANGUAGE": [
        "python", "java", "javascript", "typescript", "c", "c++", "c#",
        "go", "golang", "rust", "ruby", "php", "swift", "kotlin", "scala",
        "r", "matlab", "perl", "objective-c", "dart", "lua", "haskell",
        "elixir", "erlang", "clojure", "groovy", "shell", "bash", "powershell",
        "vba", "sql", "pl/sql", "cobol", "fortran", "assembly",
    ],
    "FRAMEWORK": [
        "django", "flask", "fastapi", "spring", "spring boot", "express",
        "express.js", "react", "react.js", "angular", "vue", "vue.js",
        "next.js", "nuxt.js", "svelte", "laravel", "symfony", "rails",
        "ruby on rails", ".net", "asp.net", "asp.net core", "node.js",
        "nestjs", "ember.js", "backbone.js", "jquery", "bootstrap",
        "tailwind", "tailwind css", "redux", "graphql apollo", "gin",
        "fiber", "hibernate", "struts",
    ],
    "DATABASE": [
        "mysql", "postgresql", "postgres", "mongodb", "sqlite", "oracle",
        "sql server", "mariadb", "redis", "cassandra", "dynamodb",
        "elasticsearch", "couchdb", "neo4j", "firestore", "firebase",
        "snowflake", "bigquery", "redshift", "cockroachdb", "supabase",
    ],
    "CLOUD_PLATFORM": [
        "aws", "amazon web services", "azure", "microsoft azure",
        "gcp", "google cloud platform", "google cloud", "heroku",
        "digitalocean", "ibm cloud", "oracle cloud", "alibaba cloud",
        "cloudflare", "vercel", "netlify", "openstack",
    ],
    "DEVOPS": [
        "docker", "kubernetes", "k8s", "jenkins", "gitlab ci", "github actions",
        "circleci", "travis ci", "ansible", "terraform", "puppet", "chef",
        "helm", "prometheus", "grafana", "istio", "argo cd", "argocd",
        "nagios", "splunk", "elk stack", "packer", "vagrant",
    ],
    "AI_ML": [
        "machine learning", "deep learning", "tensorflow", "pytorch", "keras",
        "scikit-learn", "sklearn", "opencv", "nlp", "natural language processing",
        "computer vision", "reinforcement learning", "xgboost", "lightgbm",
        "huggingface", "transformers", "spacy", "gliner", "llm", "generative ai",
        "neural networks", "mlops", "cnn", "rnn", "lstm", "gan", "bert", "gpt",
    ],
    "DATA_SCIENCE": [
        "pandas", "numpy", "scipy", "matplotlib", "seaborn", "plotly",
        "power bi", "tableau", "excel", "data visualization", "data analysis",
        "statistics", "a/b testing", "etl", "data engineering", "airflow",
        "spark", "apache spark", "hadoop", "kafka", "dbt", "looker",
    ],
    "MOBILE_DEVELOPMENT": [
        "android", "ios", "flutter", "react native", "xamarin", "swiftui",
        "jetpack compose", "kotlin multiplatform", "ionic", "cordova",
    ],
    "FRONTEND": [
        "html", "html5", "css", "css3", "sass", "scss", "less", "webpack",
        "vite", "babel", "figma", "sketch", "responsive design", "pwa",
        "web accessibility", "material ui", "chakra ui",
    ],
    "BACKEND": [
        "rest api", "restful api", "graphql", "grpc", "microservices",
        "websocket", "soap", "api gateway", "message queue", "rabbitmq",
        "celery", "nginx", "apache", "load balancing",
    ],
    "TESTING": [
        "unit testing", "pytest", "junit", "selenium", "cypress", "jest",
        "mocha", "chai", "testng", "postman", "jmeter", "cucumber",
        "test automation", "tdd", "bdd", "qa", "quality assurance",
    ],
    "CYBERSECURITY": [
        "penetration testing", "ethical hacking", "owasp", "burp suite",
        "wireshark", "metasploit", "siem", "vulnerability assessment",
        "network security", "cryptography", "iso 27001", "nist",
        "incident response", "firewall management",
    ],
    "ERP": [
        "sap", "sap abap", "sap fico", "sap mm", "oracle ebs", "netsuite",
        "workday", "peoplesoft", "microsoft dynamics", "salesforce",
    ],
    "DESIGN_TOOLS": [
        "photoshop", "illustrator", "adobe xd", "indesign", "after effects",
        "premiere pro", "canva", "invision", "framer", "blender",
    ],
    "PRODUCTIVITY_TOOLS": [
        "microsoft office", "ms office", "word", "powerpoint", "outlook",
        "google workspace", "jira", "confluence", "trello", "asana",
        "notion", "slack", "monday.com", "clickup",
    ],
    "OPERATING_SYSTEM": [
        "linux", "unix", "windows", "macos", "ubuntu", "centos", "red hat",
        "debian", "fedora",
    ],
    "TOOL": [
        "git", "github", "gitlab", "bitbucket", "svn", "vs code",
        "visual studio", "intellij", "pycharm", "eclipse", "postman",
        "docker compose", "make", "cmake",
    ],
}

# Flat name -> category lookup for fast matching
_SKILL_NAME_TO_CATEGORY: Dict[str, str] = {}
for _cat, _names in SKILL_DICTIONARY.items():
    for _n in _names:
        _SKILL_NAME_TO_CATEGORY[_n.lower()] = _cat

_ALL_SKILL_NAMES = sorted(_SKILL_NAME_TO_CATEGORY.keys(), key=len, reverse=True)

_CATEGORY_TO_MODEL_CHOICE = {
    "PROGRAMMING_LANGUAGE": SkillCategory.LANGUAGE,
    "FRAMEWORK": SkillCategory.FRAMEWORK,
    "DATABASE": SkillCategory.DATABASE,
    "CLOUD_PLATFORM": SkillCategory.CLOUD,
    "DEVOPS": SkillCategory.TOOL,
    "AI_ML": SkillCategory.TECHNICAL,
    "DATA_SCIENCE": SkillCategory.TECHNICAL,
    "MOBILE_DEVELOPMENT": SkillCategory.TECHNICAL,
    "FRONTEND": SkillCategory.TECHNICAL,
    "BACKEND": SkillCategory.TECHNICAL,
    "TESTING": SkillCategory.TOOL,
    "CYBERSECURITY": SkillCategory.TECHNICAL,
    "ERP": SkillCategory.TOOL,
    "DESIGN_TOOLS": SkillCategory.TOOL,
    "PRODUCTIVITY_TOOLS": SkillCategory.TOOL,
    "OPERATING_SYSTEM": SkillCategory.OTHER,
    "TOOL": SkillCategory.TOOL,
    "SOFT_SKILL": SkillCategory.SOFT,
    "UNKNOWN": SkillCategory.OTHER,
}

SOFT_SKILL_KEYWORDS = [
    "communication", "leadership", "teamwork", "problem solving",
    "problem-solving", "time management", "adaptability", "creativity",
    "critical thinking", "collaboration", "negotiation", "public speaking",
    "conflict resolution", "decision making", "emotional intelligence",
    "mentoring", "presentation skills", "analytical skills",
]


# ======================================================================
# Text extraction (PDF via PyMuPDF, DOCX via python-docx)
# ======================================================================

def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from a PDF using PyMuPDF only (no OCR)."""
    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text (paragraphs + tables) from a .docx file."""
    document = DocxDocument(file_path)
    lines: List[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def extract_text(file_path: str) -> str:
   
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    
    if suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported resume file type: {suffix}")


def normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def get_lines(text: str) -> List[str]:
    return [ln.strip() for ln in text.split("\n") if ln.strip()]


def strip_bullet(line: str) -> str:
    return BULLET_PREFIX_RE.sub("", line).strip()


# ======================================================================
# Section detection (RapidFuzz heading matching)
# ======================================================================

def _looks_like_heading(line: str) -> bool:
    if len(line) > 60:
        return False
    word_count = len(line.split())
    if word_count > 6:
        return False
    letters = [c for c in line if c.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    return upper_ratio > 0.5 or line.endswith(":") or word_count <= 4


def detect_sections(text: str, fuzz_threshold: int = 78) -> Dict[str, str]:
    
    lines = text.split("\n")
    sections: Dict[str, List[str]] = defaultdict(list)
    current_section: Optional[str] = None
    preamble: List[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        candidate_heading = line.lower().strip(":").strip()

        matched_section = None
        if _looks_like_heading(line):
            match = rf_process.extractOne(
                candidate_heading, _SYNONYM_TEXTS, scorer=fuzz.token_sort_ratio
            )
            if match and match[1] >= fuzz_threshold:
                idx = _SYNONYM_TEXTS.index(match[0])
                matched_section = _ALL_SYNONYMS[idx][1]

        if matched_section:
            current_section = matched_section
            continue

        if current_section:
            sections[current_section].append(line)
        else:
            preamble.append(line)

    result = {name: "\n".join(block) for name, block in sections.items()}
    result["_preamble"] = "\n".join(preamble)
    result["_full_text"] = text
    return result


def split_entries_by_blank_line_or_date(block: str) -> List[str]:
  
    if not block.strip():
        return []

    raw_chunks = re.split(r"\n\s*\n", block)
    if len(raw_chunks) > 1:
        return [c.strip() for c in raw_chunks if c.strip()]

    lines = get_lines(block)
    entries: List[List[str]] = []
    current: List[str] = []
    for line in lines:
        starts_new = bool(DATE_TOKEN_RE.search(line)) and len(line.split()) <= 8
        is_bullet = bool(BULLET_PREFIX_RE.match(line))
        if starts_new and not is_bullet and current:
            entries.append(current)
            current = [line]
        else:
            current.append(line)
    if current:
        entries.append(current)

    return ["\n".join(e) for e in entries] if entries else [block.strip()]



def run_spacy_ner(text: str) -> Dict[str, List[str]]:
    nlp = get_spacy_model()
    result: Dict[str, List[str]] = defaultdict(list)
    if nlp is None or not text.strip():
        return result
    # spaCy has default max_length; guard very long resumes
    doc = nlp(text[: nlp.max_length - 1] if len(text) >= nlp.max_length else text)
    for ent in doc.ents:
        if ent.label_ in ("PERSON", "ORG", "GPE", "DATE"):
            result[ent.label_].append(ent.text.strip())
    return result


def run_gliner_ner(text: str, labels: Optional[Sequence[str]] = None,
                    threshold: float = 0.4) -> List[Dict[str, Any]]:
    model = get_gliner_model()
    if model is None or not text.strip():
        return []
    labels = list(labels) if labels else GLINER_LABELS
    entities: List[Dict[str, Any]] = []
    # GLiNER models have a token budget; chunk very long text defensively.
    max_chars = 3000
    for start in range(0, len(text), max_chars):
        chunk = text[start:start + max_chars]
        try:
            chunk_entities = model.predict_entities(chunk, labels, threshold=threshold)
        except Exception:  # pragma: no cover - defensive
            logger.exception("GLiNER prediction failed on a chunk")
            continue
        entities.extend(chunk_entities)
    return entities


def gliner_values_for_label(entities: List[Dict[str, Any]], label: str) -> List[str]:
    return [e["text"].strip() for e in entities if e.get("label", "").lower() == label.lower()]



def extract_email(text: str) -> Optional[str]:
    matches = EMAIL_RE.findall(text)
    if not matches:
        return None
    # Prefer the first email that appears near the top of the document.
    return matches[0].strip().rstrip(".,;")


def extract_phone_numbers(text: str) -> List[str]:
   
    found: "OrderedDict[str, None]" = OrderedDict()

    for match in phonenumbers.PhoneNumberMatcher(text, "IN"):
        formatted = phonenumbers.format_number(
            match.number, phonenumbers.PhoneNumberFormat.E164
        )
        found[formatted] = None

    if not found:
        for match in phonenumbers.PhoneNumberMatcher(text, "US"):
            formatted = phonenumbers.format_number(
                match.number, phonenumbers.PhoneNumberFormat.E164
            )
            found[formatted] = None

    if not found:
        for raw in PHONE_CANDIDATE_RE.findall(text):
            candidate = "".join(raw) if isinstance(raw, tuple) else raw
            digits = re.sub(r"\D", "", candidate)
            if 7 <= len(digits) <= 15:
                found[candidate.strip()] = None

    return list(found.keys())


def extract_urls(text: str) -> Dict[str, str]:
    result = {"linkedin_url": "", "github_url": "", "portfolio_url": "", "website_url": ""}

    linkedin = LINKEDIN_RE.search(text)
    if linkedin:
        result["linkedin_url"] = _ensure_scheme(linkedin.group(0))

    github = GITHUB_RE.search(text)
    if github:
        result["github_url"] = _ensure_scheme(github.group(0))

    all_urls = URL_RE.findall(text)
    generic_urls = [
        u for u in all_urls
        if "linkedin.com" not in u.lower()
        and "github.com" not in u.lower()
        and EMAIL_RE.fullmatch(u) is None
        and "@" not in u
    ]
    if generic_urls:
        portfolio_candidates = [
            u for u in generic_urls
            if any(k in u.lower() for k in ("portfolio", "behance", "dribbble"))
        ]
        if portfolio_candidates:
            result["portfolio_url"] = _ensure_scheme(portfolio_candidates[0])
            generic_urls.remove(portfolio_candidates[0])
        if generic_urls:
            result["website_url"] = _ensure_scheme(generic_urls[0])

    return result


def _ensure_scheme(url: str) -> str:
    url = url.strip().rstrip(".,;")
    if not url.lower().startswith(("http://", "https://")):
        url = "https://" + url
    return url


def extract_name(text: str, spacy_ents: Dict[str, List[str]],
                  gliner_entities: List[Dict[str, Any]]) -> Optional[Dict[str, str]]:
    resolver = FieldResolver("full_name")

    header_lines = get_lines(text)[:5]
    for i, line in enumerate(header_lines):
        if EMAIL_RE.search(line) or URL_RE.search(line) or any(ch.isdigit() for ch in line):
            continue
        words = line.split()
        if 1 < len(words) <= 4 and all(w[:1].isupper() for w in words if w[:1].isalpha()):
            resolver.add(line, confidence=0.6 + (0.1 * (5 - i) / 5), source="heuristic_header")
            break

    gliner_names = gliner_values_for_label(gliner_entities, "Full Name")
    if gliner_names:
        resolver.add(gliner_names[0], confidence=0.85, source="gliner")

    spacy_persons = spacy_ents.get("PERSON", [])
    if spacy_persons:
        resolver.add(spacy_persons[0], confidence=0.75, source="spacy")

    best = resolver.value
    if not best:
        return None

    parsed = HumanName(best)
    return {
        "full_name": best.strip(),
        "first_name": parsed.first.strip(),
        "last_name": (parsed.last or parsed.middle).strip(),
    }


COUNTRY_NAMES = {c.name.lower(): c.name for c in pycountry.countries}
COUNTRY_ALPHA2 = {c.alpha_2.lower(): c.name for c in pycountry.countries}


def extract_country(text: str) -> Optional[str]:
    lowered = text.lower()
    for name_lower, canonical in COUNTRY_NAMES.items():
        if re.search(r"\b" + re.escape(name_lower) + r"\b", lowered):
            return canonical
    common_aliases = {
        "usa": "United States", "us": "United States", "uk": "United Kingdom",
        "uae": "United Arab Emirates",
    }
    for alias, canonical in common_aliases.items():
        if re.search(r"\b" + re.escape(alias) + r"\b", lowered):
            return canonical
    return None


def extract_location(text: str, spacy_ents: Dict[str, List[str]],
                      gliner_entities: List[Dict[str, Any]]) -> Dict[str, str]:
    result = {"address": "", "city": "", "state": "", "country": "", "postal_code": ""}

    header_block = "\n".join(get_lines(text)[:8])

    gliner_addresses = gliner_values_for_label(gliner_entities, "Address")
    if gliner_addresses:
        result["address"] = gliner_addresses[0]

    gpe_candidates = spacy_ents.get("GPE", [])
    country = extract_country(header_block) or extract_country(text)
    if country:
        result["country"] = country

    remaining_gpe = [g for g in gpe_candidates if g.lower() != (country or "").lower()]
    if remaining_gpe:
        result["city"] = remaining_gpe[0]
        if len(remaining_gpe) > 1:
            result["state"] = remaining_gpe[1]

    postal = POSTAL_CODE_RE.search(header_block)
    if postal:
        result["postal_code"] = postal.group(0)

    if not result["address"] and remaining_gpe:
        result["address"] = ", ".join(remaining_gpe)

    return result




def parse_date_string(raw: str) -> Optional[date]:
    """Hybrid date parsing: dateutil first, dateparser as a fallback."""
    if not raw:
        return None
    raw = raw.strip().strip(",")
    if PRESENT_RE.match(raw):
        return None  # caller interprets None + is_current separately

    try:
        parsed = dateutil_parser.parse(raw, default=datetime(1900, 1, 1), fuzzy=True)
        if parsed.year != 1900 or re.search(r"\d{4}", raw):
            return parsed.date()
    except (ValueError, OverflowError):
        pass

    if HAS_DATEPARSER:
        parsed = dateparser.parse(raw)
        if parsed:
            return parsed.date()

    year_match = re.search(r"(19|20)\d{2}", raw)
    if year_match:
        try:
            return date(int(year_match.group(0)), 1, 1)
        except ValueError:
            return None
    return None


def extract_date_range(text_segment: str) -> Tuple[Optional[date], Optional[date], bool]:
    """
    Finds the first date-range-looking substring in text_segment and
    returns (start_date, end_date, is_current).
    """
    tokens = DATE_TOKEN_RE.findall(text_segment)
    is_current = bool(PRESENT_RE.search(text_segment)) or bool(
        re.search(r"present|current|ongoing|till date", text_segment, re.I)
    )

    range_match = re.search(
        DATE_TOKEN_RE.pattern + DATE_RANGE_SEP_RE.pattern +
        r"(?:" + DATE_TOKEN_RE.pattern + r"|present|current|till date|ongoing|now)",
        text_segment,
        re.IGNORECASE,
    )

    start_date = end_date = None
    if range_match:
        chunk = range_match.group(0)
        parts = DATE_RANGE_SEP_RE.split(chunk, maxsplit=1)
        if len(parts) == 2:
            start_date = parse_date_string(parts[0])
            if PRESENT_RE.match(parts[1].strip()):
                is_current = True
            else:
                end_date = parse_date_string(parts[1])
    elif tokens:
        start_date = parse_date_string(tokens[0])
        if len(tokens) > 1:
            end_date = parse_date_string(tokens[1])

    return start_date, end_date, is_current


# ---------------------------------------------------------------- skills

def extract_skills(full_text: str, skills_section_text: str,
                    gliner_entities: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    found: "OrderedDict[str, Dict[str, Any]]" = OrderedDict()

    search_space = (skills_section_text or "") + "\n" + full_text
    lowered = search_space.lower()

    for name in _ALL_SKILL_NAMES:
        pattern = r"(?<![a-z0-9])" + re.escape(name) + r"(?![a-z0-9])"
        if re.search(pattern, lowered):
            category = _SKILL_NAME_TO_CATEGORY[name]
            key = name.lower()
            if key not in found:
                found[key] = {
                    "skill_name": name.title() if name.islower() and len(name) > 3 else name,
                    "category": _CATEGORY_TO_MODEL_CHOICE.get(category, SkillCategory.OTHER),
                    "confidence": 0.9,
                    "source": "dictionary",
                }

    for soft in SOFT_SKILL_KEYWORDS:
        if re.search(r"\b" + re.escape(soft) + r"\b", lowered):
            key = soft.lower()
            if key not in found:
                found[key] = {
                    "skill_name": soft.title(),
                    "category": SkillCategory.SOFT,
                    "confidence": 0.7,
                    "source": "dictionary",
                }

    gliner_skill_labels = ["Skill", "Programming Language", "Framework", "Database",
                            "Cloud Platform", "Tool", "Technology"]
    for label in gliner_skill_labels:
        for value in gliner_values_for_label(gliner_entities, label):
            key = value.lower().strip()
            if not key or key in found:
                continue
            category = _SKILL_NAME_TO_CATEGORY.get(key, "UNKNOWN")
            found[key] = {
                "skill_name": value.strip(),
                "category": _CATEGORY_TO_MODEL_CHOICE.get(category, SkillCategory.OTHER),
                "confidence": 0.6,
                "source": "gliner",
            }

    if skills_section_text:
        for raw_line in skills_section_text.split("\n"):
            line = strip_bullet(raw_line)
            if ":" in line:
                line = line.split(":", 1)[1]
            for token in re.split(r"[,/|•]", line):
                token = token.strip()
                if 1 < len(token) <= 40 and not token.lower() in found:
                    category = _SKILL_NAME_TO_CATEGORY.get(token.lower(), "UNKNOWN")
                    found[token.lower()] = {
                        "skill_name": token,
                        "category": _CATEGORY_TO_MODEL_CHOICE.get(category, SkillCategory.OTHER),
                        "confidence": 0.5,
                        "source": "section_keyword",
                    }

    return list(found.values())



def extract_education(section_text: str,
                       gliner_entities: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    entries = split_entries_by_blank_line_or_date(section_text)
    results = []

    degrees_g = gliner_values_for_label(gliner_entities, "Degree")
    colleges_g = gliner_values_for_label(gliner_entities, "College")
    universities_g = gliner_values_for_label(gliner_entities, "University")

    for idx, entry in enumerate(entries):
        lines = get_lines(entry)
        if not lines:
            continue
        entry_lower = entry.lower()

        degree = None
        for kw in DEGREE_KEYWORDS:
            if kw in entry_lower:
                for line in lines:
                    if kw in line.lower():
                        degree = strip_bullet(line)
                        break
                break
        if not degree and idx < len(degrees_g):
            degree = degrees_g[idx]
        if not degree:
            degree = lines[0]

        institution = None
        if idx < len(universities_g):
            institution = universities_g[idx]
        elif idx < len(colleges_g):
            institution = colleges_g[idx]
        else:
            for line in lines[1:3]:
                if line != degree:
                    institution = strip_bullet(line)
                    break
        if not institution:
            institution = lines[1] if len(lines) > 1 else lines[0]

        start_date, end_date, is_current = extract_date_range(entry)

        cgpa = None
        cgpa_match = CGPA_RE.search(entry)
        if cgpa_match:
            try:
                cgpa = Decimal(cgpa_match.group(1))
            except InvalidOperation:
                cgpa = None

        percentage = None
        pct_match = PERCENTAGE_RE.search(entry)
        if pct_match:
            try:
                percentage = Decimal(pct_match.group(1))
            except InvalidOperation:
                percentage = None

        results.append({
            "degree": degree[:255],
            "field_of_study": "",
            "institution_name": institution[:255],
            "university": institution[:255],
            "start_date": start_date,
            "end_date": end_date,
            "is_current": is_current,
            "cgpa": cgpa,
            "percentage": percentage,
            "description": entry.strip()[:2000],
            "display_order": idx + 1,
        })

    return results




EMPLOYMENT_TYPE_KEYWORDS = {
    EmploymentType.INTERNSHIP: ["intern", "internship"],
    EmploymentType.PART_TIME: ["part-time", "part time"],
    EmploymentType.CONTRACT: ["contract", "contractor"],
    EmploymentType.FREELANCE: ["freelance", "freelancer"],
    EmploymentType.TEMPORARY: ["temporary", "temp"],
    EmploymentType.APPRENTICESHIP: ["apprentice", "apprenticeship"],
    EmploymentType.FULL_TIME: ["full-time", "full time"],
}


def _guess_employment_type(text_block: str) -> str:
    lowered = text_block.lower()
    for choice, keywords in EMPLOYMENT_TYPE_KEYWORDS.items():
        if any(kw in lowered for kw in keywords):
            return choice
    return ""


def extract_experience(section_text: str,
                        gliner_entities: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    entries = split_entries_by_blank_line_or_date(section_text)
    results = []

    companies_g = gliner_values_for_label(gliner_entities, "Company")
    titles_g = gliner_values_for_label(gliner_entities, "Job Title")
    responsibilities_g = gliner_values_for_label(gliner_entities, "Responsibility")

    for idx, entry in enumerate(entries):
        lines = get_lines(entry)
        if not lines:
            continue

        title = titles_g[idx] if idx < len(titles_g) else lines[0]
        company = companies_g[idx] if idx < len(companies_g) else (
            lines[1] if len(lines) > 1 else ""
        )

        start_date, end_date, is_current = extract_date_range(entry)

        bullet_lines = [strip_bullet(l) for l in lines if BULLET_PREFIX_RE.match(l)]
        if not bullet_lines:
            bullet_lines = [l for l in lines[2:] if len(l.split()) > 3]

        responsibilities = bullet_lines if bullet_lines else responsibilities_g[:5]

        achievements = [
            b for b in bullet_lines
            if re.search(r"\b(increased|reduced|improved|achieved|awarded|won|delivered|"
                         r"grew|saved|launched|led)\b", b, re.I)
        ]

        technologies_found = []
        entry_lower = entry.lower()
        for name in _ALL_SKILL_NAMES:
            if re.search(r"(?<![a-z0-9])" + re.escape(name) + r"(?![a-z0-9])", entry_lower):
                technologies_found.append(name)

        results.append({
            "company_name": company.strip()[:255] or "Unknown Company",
            "designation": title.strip()[:255] or "Unknown Title",
            "employment_type": _guess_employment_type(entry),
            "location": "",
            "start_date": start_date,
            "end_date": end_date,
            "is_current": is_current,
            "responsibilities": responsibilities,
            "achievements": achievements,
            "technologies": technologies_found,
            "skills_used": technologies_found,
            "description": entry.strip()[:3000],
            "display_order": idx + 1,
        })

    return results




def extract_projects(section_text: str,
                      gliner_entities: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    entries = split_entries_by_blank_line_or_date(section_text)
    results = []

    titles_g = gliner_values_for_label(gliner_entities, "Project")

    for idx, entry in enumerate(entries):
        lines = get_lines(entry)
        if not lines:
            continue

        title = titles_g[idx] if idx < len(titles_g) else strip_bullet(lines[0])

        urls = extract_urls(entry)
        start_date, end_date, _ = extract_date_range(entry)

        bullet_lines = [strip_bullet(l) for l in lines if BULLET_PREFIX_RE.match(l)]

        technologies_found = []
        entry_lower = entry.lower()
        for name in _ALL_SKILL_NAMES:
            if re.search(r"(?<![a-z0-9])" + re.escape(name) + r"(?![a-z0-9])", entry_lower):
                technologies_found.append(name)

        results.append({
            "project_title": title.strip()[:255],
            "role": "",
            "organization": "",
            "technologies": technologies_found,
            "description": entry.strip()[:3000],
            "responsibilities": bullet_lines,
            "project_url": urls.get("website_url", ""),
            "github_url": urls.get("github_url", ""),
            "start_date": start_date,
            "end_date": end_date,
            "display_order": idx + 1,
        })

    return results




def extract_certifications(section_text: str,
                            gliner_entities: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    entries = split_entries_by_blank_line_or_date(section_text)
    results = []

    certs_g = gliner_values_for_label(gliner_entities, "Certification")

    for idx, entry in enumerate(entries):
        lines = get_lines(entry)
        if not lines:
            continue

        name = certs_g[idx] if idx < len(certs_g) else strip_bullet(lines[0])
        issuer = ""
        if " by " in name.lower():
            parts = re.split(r"\s+by\s+", name, flags=re.I, maxsplit=1)
            name, issuer = parts[0].strip(), parts[1].strip()
        elif len(lines) > 1:
            issuer = strip_bullet(lines[1])

        issue_date, expiry_date, _ = extract_date_range(entry)

        cred_id_match = re.search(r"(?:credential\s*id|id)\s*[:#]?\s*([\w\-]+)", entry, re.I)
        credential_id = cred_id_match.group(1) if cred_id_match else ""

        url_match = URL_RE.search(entry)
        credential_url = _ensure_scheme(url_match.group(0)) if url_match else ""

        results.append({
            "certification_name": name.strip()[:255],
            "issuing_organization": issuer.strip()[:255],
            "issue_date": issue_date,
            "expiry_date": expiry_date,
            "credential_id": credential_id[:255],
            "credential_url": credential_url,
            "display_order": idx + 1,
        })

    return results




LANGUAGE_NAMES = [
    "english", "hindi", "tamil", "telugu", "kannada", "malayalam", "marathi",
    "gujarati", "punjabi", "bengali", "urdu", "spanish", "french", "german",
    "mandarin", "chinese", "japanese", "korean", "italian", "portuguese",
    "russian", "arabic", "dutch", "swedish", "turkish", "vietnamese",
    "thai", "polish", "greek",
]

LANGUAGE_PROFICIENCY_KEYWORDS = {
    LanguageProficiency.NATIVE: ["native", "mother tongue"],
    LanguageProficiency.FLUENT: ["fluent", "professional working proficiency", "advanced"],
    LanguageProficiency.INTERMEDIATE: ["intermediate", "conversational"],
    LanguageProficiency.BASIC: ["basic", "beginner", "elementary"],
}


def extract_languages(section_text: str) -> List[Dict[str, Any]]:
    if not section_text.strip():
        return []
    results = []
    lowered = section_text.lower()

    for idx, lang in enumerate(LANGUAGE_NAMES):
        pattern = r"\b" + re.escape(lang) + r"\b"
        match = re.search(pattern, lowered)
        if not match:
            continue

        window = lowered[match.end(): match.end() + 40]
        proficiency = LanguageProficiency.BASIC
        for level, keywords in LANGUAGE_PROFICIENCY_KEYWORDS.items():
            if any(kw in window for kw in keywords):
                proficiency = level
                break

        results.append({
            "language": lang.title(),
            "proficiency": proficiency,
            "can_read": True,
            "can_write": True,
            "can_speak": True,
            "display_order": len(results) + 1,
        })

    return results



def extract_achievements(section_text: str,
                          gliner_entities: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    if not section_text.strip():
        return []

    achievements_g = gliner_values_for_label(gliner_entities, "Achievement")
    entries = split_entries_by_blank_line_or_date(section_text)
    results = []

    for idx, entry in enumerate(entries):
        lines = get_lines(entry)
        if not lines:
            continue
        title = achievements_g[idx] if idx < len(achievements_g) else strip_bullet(lines[0])
        achievement_date, _, _ = extract_date_range(entry)

        results.append({
            "title": title.strip()[:255],
            "organization": "",
            "achievement_date": achievement_date,
            "description": entry.strip()[:2000],
            "display_order": idx + 1,
        })

    return results



def extract_summary(sections: Dict[str, str]) -> str:
    summary = sections.get("summary", "").strip()
    if summary:
        return summary[:3000]
    preamble = sections.get("_preamble", "")
    candidate_lines = [
        l for l in get_lines(preamble)
        if not EMAIL_RE.search(l) and not URL_RE.search(l)
        and len(l.split()) > 6
    ]
    return " ".join(candidate_lines)[:3000]




def parse_resume_file(file_path: str) -> Dict[str, Any]:
    
    parser_errors: List[Dict[str, str]] = []
    parsed: Dict[str, Any] = {
        "raw_text": "",
        "personal_information": {},
        "education": [],
        "experience": [],
        "skills": [],
        "projects": [],
        "certifications": [],
        "languages": [],
        "achievements": [],
        "parser_errors": parser_errors,
    }

    def safe_run(section_name: str, func, *args, default):
        try:
            return func(*args)
        except Exception as exc:  # noqa: BLE001 - intentional broad catch
            logger.exception("Failed to parse section '%s'", section_name)
            parser_errors.append({"section": section_name, "error": str(exc)})
            return default

    raw_text = safe_run("text_extraction", extract_text, file_path, default="")
    text = normalize_whitespace(raw_text)
    parsed["raw_text"] = text

    if not text.strip():
        parser_errors.append({"section": "text_extraction", "error": "No text could be extracted."})
        return parsed

    sections = safe_run("section_detection", detect_sections, text, default={"_full_text": text})

    spacy_ents = safe_run("spacy_ner", run_spacy_ner, text, default={})
    gliner_entities = safe_run(
        "gliner_ner", run_gliner_ner, text, GLINER_LABELS, default=[]
    )

    # ---- personal information -------------------------------------
    def build_personal_information():
        name_info = extract_name(text, spacy_ents, gliner_entities) or {}
        location_info = extract_location(text, spacy_ents, gliner_entities)
        urls = extract_urls(text)
        phones = extract_phone_numbers(text)
        email = extract_email(text)
        summary = extract_summary(sections)

        return {
            "first_name": name_info.get("first_name", ""),
            "last_name": name_info.get("last_name", ""),
            "full_name": name_info.get("full_name", ""),
            "email": email or "",
            "phone_number": phones[0] if phones else "",
            "alternate_phone_number": phones[1] if len(phones) > 1 else "",
            "profile_summary": summary,
            "nationality": "",
            "address": location_info.get("address", ""),
            "city": location_info.get("city", ""),
            "state": location_info.get("state", ""),
            "country": location_info.get("country", ""),
            "postal_code": location_info.get("postal_code", ""),
            "linkedin_url": urls.get("linkedin_url", ""),
            "github_url": urls.get("github_url", ""),
            "portfolio_url": urls.get("portfolio_url", ""),
            "website_url": urls.get("website_url", ""),
        }

    parsed["personal_information"] = safe_run(
        "personal_information", build_personal_information, default={}
    )

    parsed["education"] = safe_run(
        "education", extract_education, sections.get("education", ""), gliner_entities,
        default=[],
    )
    parsed["experience"] = safe_run(
        "experience", extract_experience, sections.get("experience", ""), gliner_entities,
        default=[],
    )
    parsed["skills"] = safe_run(
        "skills", extract_skills, text, sections.get("skills", ""), gliner_entities,
        default=[],
    )
    parsed["projects"] = safe_run(
        "projects", extract_projects, sections.get("projects", ""), gliner_entities,
        default=[],
    )
    parsed["certifications"] = safe_run(
        "certifications", extract_certifications, sections.get("certifications", ""),
        gliner_entities, default=[],
    )
    parsed["languages"] = safe_run(
        "languages", extract_languages, sections.get("languages", ""), default=[],
    )
    parsed["achievements"] = safe_run(
        "achievements", extract_achievements, sections.get("achievements", ""),
        gliner_entities, default=[],
    )

    return parsed



def _to_employment_type_choice(value: str) -> str:
    valid_values = {choice for choice, _ in EmploymentType.choices}
    return value if value in valid_values else ""


def _to_skill_category_choice(value: str) -> str:
    valid_values = {choice for choice, _ in SkillCategory.choices}
    return value if value in valid_values else SkillCategory.OTHER


def _to_language_proficiency_choice(value: str) -> str:
    valid_values = {choice for choice, _ in LanguageProficiency.choices}
    return value if value in valid_values else LanguageProficiency.BASIC


@transaction.atomic
def save_parsed_resume_to_db(resume: "Resume", parsed: Dict[str, Any]) -> Dict[str, Any]:
   
    parser_errors: List[Dict[str, str]] = list(resume.parser_errors or [])
    parser_errors.extend(parsed.get("parser_errors", []))

    def safe_save(section_name: str, func):
        try:
            func()
        except Exception as exc:  # noqa: BLE001 - intentional broad catch
            logger.exception("Failed to save section '%s' for resume %s", section_name, resume.pk)
            parser_errors.append({"section": section_name, "error": str(exc)})

   
    def save_personal_information():
        info = parsed.get("personal_information") or {}
        if not info:
            return
        ResumePersonalInformation.objects.update_or_create(
            resume=resume,
            defaults={
                "first_name": info.get("first_name", "")[:100],
                "last_name": info.get("last_name", "")[:100],
                "full_name": info.get("full_name", "")[:255],
                "email": info.get("email", "")[:254],
                "phone_number": info.get("phone_number", "")[:20],
                "alternate_phone_number": info.get("alternate_phone_number", "")[:20],
                "profile_summary": info.get("profile_summary", ""),
                "nationality": info.get("nationality", "")[:100],
                "address": info.get("address", ""),
                "city": info.get("city", "")[:100],
                "state": info.get("state", "")[:100],
                "country": info.get("country", "")[:100],
                "postal_code": info.get("postal_code", "")[:20],
                "linkedin_url": info.get("linkedin_url", "")[:200],
                "github_url": info.get("github_url", "")[:200],
                "portfolio_url": info.get("portfolio_url", "")[:200],
                "website_url": info.get("website_url", "")[:200],
            },
        )

    def save_education():
        entries = parsed.get("education") or []
        ResumeEducation.objects.filter(resume=resume).delete()
        objs = [
            ResumeEducation(
                resume=resume,
                degree=e.get("degree", "")[:255] or "Unknown Degree",
                field_of_study=e.get("field_of_study", "")[:255],
                institution_name=e.get("institution_name", "")[:255] or "Unknown Institution",
                university=e.get("university", "")[:255],
                start_date=e.get("start_date"),
                end_date=e.get("end_date"),
                is_current=bool(e.get("is_current", False)),
                cgpa=e.get("cgpa"),
                percentage=e.get("percentage"),
                description=e.get("description", ""),
                display_order=e.get("display_order", i + 1),
            )
            for i, e in enumerate(entries)
        ]
        if objs:
            ResumeEducation.objects.bulk_create(objs)

    def save_experience():
        entries = parsed.get("experience") or []
        ResumeWorkExperience.objects.filter(resume=resume).delete()
        objs = [
            ResumeWorkExperience(
                resume=resume,
                company_name=e.get("company_name", "")[:255] or "Unknown Company",
                designation=e.get("designation", "")[:255] or "Unknown Title",
                employment_type=_to_employment_type_choice(e.get("employment_type", "")),
                location=e.get("location", "")[:255],
                start_date=e.get("start_date"),
                end_date=e.get("end_date"),
                is_current=bool(e.get("is_current", False)),
                responsibilities=e.get("responsibilities", []),
                achievements=e.get("achievements", []),
                technologies=e.get("technologies", []),
                skills_used=e.get("skills_used", []),
                description=e.get("description", ""),
                display_order=e.get("display_order", i + 1),
            )
            for i, e in enumerate(entries)
        ]
        if objs:
            ResumeWorkExperience.objects.bulk_create(objs)

    def save_skills():
        entries = parsed.get("skills") or []
        ResumeSkill.objects.filter(resume=resume).delete()
        seen = set()
        objs = []
        for i, s in enumerate(entries):
            name = (s.get("skill_name") or "").strip()[:150]
            if not name or name.lower() in seen:
                continue
            seen.add(name.lower())
            objs.append(
                ResumeSkill(
                    resume=resume,
                    skill_name=name,
                    category=_to_skill_category_choice(s.get("category", SkillCategory.OTHER)),
                    proficiency="",
                    years_of_experience=None,
                    last_used=None,
                    display_order=i + 1,
                )
            )
        if objs:
            ResumeSkill.objects.bulk_create(objs, ignore_conflicts=True)

    def save_projects():
        entries = parsed.get("projects") or []
        ResumeProject.objects.filter(resume=resume).delete()
        objs = [
            ResumeProject(
                resume=resume,
                project_title=p.get("project_title", "")[:255] or "Untitled Project",
                role=p.get("role", "")[:255],
                organization=p.get("organization", "")[:255],
                technologies=p.get("technologies", []),
                description=p.get("description", ""),
                responsibilities=p.get("responsibilities", []),
                project_url=p.get("project_url", "")[:200],
                github_url=p.get("github_url", "")[:200],
                start_date=p.get("start_date"),
                end_date=p.get("end_date"),
                display_order=p.get("display_order", i + 1),
            )
            for i, p in enumerate(entries)
        ]
        if objs:
            ResumeProject.objects.bulk_create(objs)

    def save_certifications():
        entries = parsed.get("certifications") or []
        ResumeCertification.objects.filter(resume=resume).delete()
        objs = [
            ResumeCertification(
                resume=resume,
                certification_name=c.get("certification_name", "")[:255] or "Unknown Certification",
                issuing_organization=c.get("issuing_organization", "")[:255],
                issue_date=c.get("issue_date"),
                expiry_date=c.get("expiry_date"),
                credential_id=c.get("credential_id", "")[:255],
                credential_url=c.get("credential_url", "")[:200],
                display_order=c.get("display_order", i + 1),
            )
            for i, c in enumerate(entries)
        ]
        if objs:
            ResumeCertification.objects.bulk_create(objs)

    def save_languages():
        entries = parsed.get("languages") or []
        ResumeLanguage.objects.filter(resume=resume).delete()
        objs = [
            ResumeLanguage(
                resume=resume,
                language=l.get("language", "")[:100] or "Unknown",
                proficiency=_to_language_proficiency_choice(l.get("proficiency", "")),
                can_read=bool(l.get("can_read", True)),
                can_write=bool(l.get("can_write", True)),
                can_speak=bool(l.get("can_speak", True)),
                display_order=l.get("display_order", i + 1),
            )
            for i, l in enumerate(entries)
        ]
        if objs:
            ResumeLanguage.objects.bulk_create(objs)

    def save_achievements():
        entries = parsed.get("achievements") or []
        ResumeAchievement.objects.filter(resume=resume).delete()
        objs = [
            ResumeAchievement(
                resume=resume,
                title=a.get("title", "")[:255] or "Untitled Achievement",
                organization=a.get("organization", "")[:255],
                achievement_date=a.get("achievement_date"),
                description=a.get("description", ""),
                display_order=a.get("display_order", i + 1),
            )
            for i, a in enumerate(entries)
        ]
        if objs:
            ResumeAchievement.objects.bulk_create(objs)

    safe_save("personal_information", save_personal_information)
    safe_save("education", save_education)
    safe_save("experience", save_experience)
    safe_save("skills", save_skills)
    safe_save("projects", save_projects)
    safe_save("certifications", save_certifications)
    safe_save("languages", save_languages)
    safe_save("achievements", save_achievements)

    resume.extracted_text = parsed.get("raw_text", "")[:1_000_000]
    resume.parser_errors = parser_errors
    resume.status = ResumeStatus.FAILED if parser_errors and not _has_any_saved_data(resume) \
        else ResumeStatus.COMPLETED
    resume.save(update_fields=["extracted_text", "parser_errors", "status", "updated_at"])

    return {"parser_errors": parser_errors, "status": resume.status}


def _has_any_saved_data(resume: "Resume") -> bool:
    return (
        ResumePersonalInformation.objects.filter(resume=resume).exists()
        or resume.educations.exists()
        or resume.work_experiences.exists()
        or resume.skills.exists()
    )


def process_resume(resume: "Resume") -> "Resume":
   
    resume.status = ResumeStatus.PROCESSING
    resume.save(update_fields=["status", "updated_at"])

    try:
        file_path = resume.resume_file.path
    except Exception as exc:  
        resume.status = ResumeStatus.FAILED
        resume.parser_errors = list(resume.parser_errors or []) + [
            {"section": "file_access", "error": str(exc)}
        ]
        resume.save(update_fields=["status", "parser_errors", "updated_at"])
        return resume

    parsed = parse_resume_file(file_path)
    save_parsed_resume_to_db(resume, parsed)

    resume.refresh_from_db()
    return resume

